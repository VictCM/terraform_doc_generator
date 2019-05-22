from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
import json
import time
from resource_parser import (dict_networks, dict_nodes)


variables_file = open('terraform_output.json')
variables = json.loads(variables_file.read())
variables_file.close()
filename='./tests/origen.docx'
document = Document()



# Reemplaza las palabras que hemos definido en el terraform_output.json
def replace_keywords():
    for p in document.paragraphs:
        inline = p.runs
        for j in range(0,len(inline)):
            for keyword, variable in variables.items():
                
                ### Keyword lo utilizamos para encontrar el recurso, por ejemplo, las instancias. ###
                #####################################################################################



                inline[j].text = inline[j].text.replace(keyword, variable)
                #print(str(keyword))
        #print(inline[j].text)

# Creamos documentación a partir de los recursos del despliegue
def add_resources_info(dict_nodes, dict_networks):
    num_nodes = len(dict_nodes)
    num_networks = len(dict_networks)

    p = document.add_paragraph()
    # Titulo
    document.add_heading('Reporte de despliegue', 0)
    # Capítulo
    document.add_heading('''Despliegue actual (Hora: ''' + time.strftime("%H:%M:%S") + ''')''', level=1)
    # Info-texto
    introduction = '''En el despliegue realizado se han creado ''' + str(num_nodes) + ''' instancias y ''' + str(num_networks) + ''' redes. \n'''
    introduction +='''  Nodos: \n'''
    p = document.add_paragraph(introduction)

    for i in dict_nodes.keys():
        p = document.add_paragraph(" ", style='List Bullet')
        r=p.add_run()
        r.add_picture('./images/deployment_graph.png',width=Inches(6))
        p.add_run(str(i) + ''': ''').bold = True

        # Busco entre las posibles opciones de despliegue que tenemos en el equipo cual es la que estamos desplegando
        found_hng=i.find("hng")
        found_pw=i.find("pw")
        found_ear=i.find("ear")
        found_uni=i.find("uni")
        num_HNG, num_vEar, num_UniManager = (0,)*3
        if (found_hng != -1) or (found_pw != -1):
            num_HNG += 1
            p.add_run('Instancia del HNG de Parallel Wireless. Gateway rural virtualizado elegido por Internet para Todos. ')
            p.add_run('Interconecta la red de acceso con el core del MNO ofreciendo independencia entre IpT y la red del MNO. ')
            p.add_run('También actúa como elemento de seguridad entre la RAN y el core, y sirve como punto de control para IpT. ')
            p.add_run('Facilita la conexión con terceras partes agregando las interfaces RAN 2G/3G/4G hacia el core. ')
            p.add_run('''Se han levantado ''' + str(num_HNG)+ ''' instancia(s), según un sabor e imagen específicas y contará con ''')
            p.add_run(str(len(dict_nodes[i])))
            p.add_run(' interfaces de red, para los distintos tipos de tráfico y necesidades. La(s) interface(s) de red tiene(n) acceso:\n')
            for j in dict_nodes[i].keys():
                p.add_run('''       - ''' + j + ''': Con IP ''' + dict_nodes[i][j] + '''.\n''')
        elif (found_ear != -1):
            num_vEar += 1
            p.add_run('Instancia del Network Ear de IpT. Solución software del equipo de Internet para Todos ')
            p.add_run('para el monitoreo y generación de alarmas dentro de la red. ')
            p.add_run('''Se han levantado ''' + str(num_vEar)+ ''' instancia(s), según un sabor e imagen específicas y contará con ''')
            p.add_run(str(len(dict_nodes[i])))
            p.add_run(' interfaces de red, con al menos una de ellas conectada a un elemento de red (HNG, en este caso). La(s) interface(s) de red tiene(n) acceso:\n')
            for j in dict_nodes[i].keys():
                p.add_run('''       - ''' + j + ''': Con IP ''' + dict_nodes[i][j] + '''.\n''')
        elif (found_uni):
            num_UniManager += 1
            p.add_run('Instancia del Unimanager. Es un Element Management System (EMS) para las antenas y HNGs de Parallel Wireless. ')
            p.add_run('Ofrece un Interfaz de Usuario Gráfica (GUI) a los operadores para el monitoreo y manejo de los elementos de red.')
            p.add_run('''Se han levantado ''' + str(num_UniManager)+ ''' instancia(s), según un sabor e imagen específicas y contará con ''')
            p.add_run(str(len(dict_nodes[i])))
            p.add_run(' interfaces de red. La(s) interface(s) de red tiene(n) acceso:\n')
            for j in dict_nodes[i].keys():
                p.add_run('''       - ''' + j + ''': Con IP ''' + dict_nodes[i][j] + '''.\n''')
        else:
            p.add_run('No se ha encontrado referencia a esa instancia')
    
   

def doc_generator(dict_nodes, dict_networks):
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture('./images/deployment_graph.png',width=Inches(6))
    #p = document.add_paragraph(" ")
    #run = p.add_run()
    #run.add_break(WD_BREAK.PAGE)
    add_resources_info(dict_nodes, dict_networks)
    replace_keywords()
    document.add_picture('./images/deployment_graph.png', width=Inches(6))
    document.save('./tests/demo.docx')